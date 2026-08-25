"""Conversion between .docx and markdown.

The orchestrator owns this conversion so the models never touch a binary format:
they receive markdown and return markdown, and this module renders the result.
A model therefore cannot corrupt a document by writing malformed OOXML.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?[\s:-]*\|[\s|:-]*$")
_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)")


class DocxError(RuntimeError):
    pass


# --- reading ------------------------------------------------------------------


def _iter_blocks(document: Document):
    """Yield paragraphs and tables in true document order."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _heading_level(style_name: str) -> int | None:
    name = (style_name or "").strip().lower()
    if name in ("title",):
        return 1
    match = re.match(r"heading (\d)", name)
    return int(match.group(1)) if match else None


def _table_to_markdown(table: Table) -> str:
    rows = [
        [cell.text.replace("\n", " ").replace("|", "\\|").strip() for cell in row.cells]
        for row in table.rows
    ]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    lines = ["| " + " | ".join(rows[0]) + " |", "|" + "---|" * width]
    lines.extend("| " + " | ".join(r) + " |" for r in rows[1:])
    return "\n".join(lines)


def extract(path: Path) -> str:
    """Read a .docx into markdown, preserving headings, lists and tables."""
    try:
        document = Document(str(path))
    except Exception as exc:
        raise DocxError(f"Could not open {path.name}: {exc}") from exc

    lines: list[str] = []
    for block in _iter_blocks(document):
        if isinstance(block, Table):
            markdown = _table_to_markdown(block)
            if markdown:
                lines.extend(["", markdown, ""])
            continue

        text = block.text.strip()
        if not text:
            continue
        style = block.style.name if block.style is not None else ""
        level = _heading_level(style)
        if level:
            lines.extend(["", "#" * level + " " + text, ""])
        elif "list bullet" in style.lower():
            lines.append(f"- {text}")
        elif "list number" in style.lower():
            lines.append(f"1. {text}")
        else:
            lines.extend([text, ""])

    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip() + "\n"


# --- writing ------------------------------------------------------------------


def _add_inline(paragraph, text: str) -> None:
    """Render **bold**, *italic* and `code` as real runs."""
    for chunk in _INLINE_RE.split(text):
        if not chunk:
            continue
        if (chunk.startswith("**") and chunk.endswith("**")) or (
            chunk.startswith("__") and chunk.endswith("__")
        ):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("`") and chunk.endswith("`") and len(chunk) > 2:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
        elif (
            chunk.startswith("*")
            and chunk.endswith("*")
            and len(chunk) > 2
        ) or (chunk.startswith("_") and chunk.endswith("_") and len(chunk) > 2):
            paragraph.add_run(chunk[1:-1]).italic = True
        else:
            paragraph.add_run(chunk)


def _split_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [c.strip().replace("\\|", "|") for c in stripped.split("|")]


def _add_table(document: Document, rows: list[list[str]]) -> None:
    width = max(len(r) for r in rows)
    table = document.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for column, value in enumerate(row + [""] * (width - len(row))):
            cells[column].text = ""
            paragraph = cells[column].paragraphs[0]
            _add_inline(paragraph, value)
            if index == 0:
                for run in paragraph.runs:
                    run.bold = True


def build(markdown: str, destination: Path, title: str | None = None) -> Path:
    """Render markdown to a new .docx. Never overwrites an existing file."""
    if destination.exists():
        raise DocxError(f"Refusing to overwrite an existing file: {destination}")

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if title:
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(16)

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            index += 1
            block: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            index += 1
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # A table is a run of pipe-delimited lines; the separator row is dropped.
        if "|" in stripped and stripped.count("|") >= 2:
            rows: list[list[str]] = []
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                if not _TABLE_SEPARATOR_RE.match(lines[index]):
                    rows.append(_split_row(lines[index]))
                index += 1
            if rows:
                _add_table(document, rows)
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            document.add_heading(heading_match.group(2).strip(), level=level)
            index += 1
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            _add_inline(document.add_paragraph(style="List Bullet"), bullet.group(1))
            index += 1
            continue

        numbered = _NUMBERED_RE.match(line)
        if numbered:
            _add_inline(document.add_paragraph(style="List Number"), numbered.group(1))
            index += 1
            continue

        if set(stripped) <= {"-", "_", "*"} and len(stripped) >= 3:
            index += 1
            continue

        _add_inline(document.add_paragraph(), stripped)
        index += 1

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    return destination


def is_readable(path: Path) -> bool:
    try:
        Document(str(path))
        return True
    except Exception:
        return False
