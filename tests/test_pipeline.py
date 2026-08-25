"""End-to-end runs with the mock provider.

Proves the parts that must be right before a real model is ever involved: naming,
phase separation, idempotency, and that the original file is never touched.
"""

from pathlib import Path

import pytest

from docx import Document

from paper_automation import phases, scanner
from paper_automation.models import FailureKind, Phase, ProcessingState, ProviderError
from paper_automation.providers.base import CliProvider
from paper_automation.storage import LocalStorage

MONTH = "August 2026"


def write_paper(folder: Path, name: str) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("A Study of Things", level=1)
    document.add_heading("Abstract", level=2)
    document.add_paragraph("We evaluate ResNet on CIFAR-10.")
    document.add_heading("Introduction", level=2)
    document.add_paragraph("Prior work is limited. " * 20)
    document.add_heading("Results", level=2)
    document.add_paragraph("Accuracy was 94.20% and F1 was 0.912.")
    document.add_heading("Conclusion", level=2)
    document.add_paragraph("It works.")
    document.add_heading("References", level=2)
    document.add_paragraph("[1] Someone. A paper. 2020.")
    path = folder / name
    document.save(str(path))
    return path


def run_pipeline(cfg, storage, phase="both"):
    cfg.provider_mode = "mock"
    month_dir = cfg.research_papers_root / MONTH
    states = scanner.scan_month(cfg, storage, month_dir)
    return phases.run(cfg, storage, MONTH, month_dir, states, phase)


def test_full_pipeline_produces_the_three_expected_files(cfg, storage):
    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    original = write_paper(folder, "Vani.docx")
    before = original.read_bytes()

    run_pipeline(cfg, storage)

    names = sorted(p.name for p in folder.iterdir() if p.is_file())
    assert names == ["Correct_Vani_paper.docx", "Vani.docx", "Vani_review.docx"]
    # Spec section 33: the original is immutable.
    assert original.read_bytes() == before


def test_second_run_changes_nothing(cfg, storage):
    """Spec section 23: safe to run repeatedly."""
    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    write_paper(folder, "Vani.docx")
    run_pipeline(cfg, storage)
    snapshot = {p.name: p.read_bytes() for p in folder.iterdir() if p.is_file()}

    run_pipeline(cfg, storage)

    after = {p.name: p.read_bytes() for p in folder.iterdir() if p.is_file()}
    assert after == snapshot


def test_no_duplicate_suffixes_are_ever_created(cfg, storage):
    """Spec section 22: never Vani_review_review.docx or Correct_Correct_....docx."""
    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    write_paper(folder, "Vani.docx")

    run_pipeline(cfg, storage)
    run_pipeline(cfg, storage)

    names = [p.name for p in folder.iterdir()]
    assert not any("_review_review" in n or n.startswith("Correct_Correct_") for n in names)
    assert len(names) == 3


def test_reviews_all_finish_before_any_revision(cfg, storage):
    """Spec section 13: the phases must not interleave."""
    order: list[tuple[str, str]] = []

    class Recording(CliProvider):
        name = model_label = "recording"

        def _invoke(self, workdir: Path, prompt: str) -> str:
            phase = "revise" if (workdir / "review.md").exists() else "review"
            order.append((phase, workdir.name))
            (workdir / "output.md").write_text(
                "# Out\n\n## Abstract\nx\n\n## Introduction\ny\n\n"
                "## Conclusion\nz\n\n## References\n[1] a\n",
                encoding="utf-8",
            )
            return "ok"

    for client in ("Vani", "R Ashok", "Jyoti R"):
        write_paper(
            cfg.research_papers_root / MONTH / "Manoj Paper" / client, f"{client}.docx"
        )

    import paper_automation.phases as phases_module

    original_builder = phases_module.build_providers
    phases_module.build_providers = lambda _cfg: (Recording(), Recording())
    try:
        run_pipeline(cfg, storage)
    finally:
        phases_module.build_providers = original_builder

    kinds = [k for k, _ in order]
    assert kinds == ["review"] * 3 + ["revise"] * 3


def test_skipped_folders_are_left_untouched(cfg, storage):
    ambiguous = cfg.research_papers_root / MONTH / "Manoj Paper" / "R Ashok"
    write_paper(ambiguous, "R_Ashok.docx")
    write_paper(ambiguous, "old.docx")
    empty = cfg.research_papers_root / MONTH / "Manoj Paper" / "Jyoti R"
    empty.mkdir(parents=True)

    run_pipeline(cfg, storage)

    assert sorted(p.name for p in ambiguous.iterdir()) == ["R_Ashok.docx", "old.docx"]
    assert list(empty.iterdir()) == []


def test_one_failure_does_not_stop_the_run(cfg, storage):
    """Spec section 24: a bad client must not end the batch."""

    class FailsOnVani(CliProvider):
        name = model_label = "flaky"

        def _invoke(self, workdir: Path, prompt: str) -> str:
            if "Vani" in workdir.name:
                raise ProviderError(FailureKind.UNKNOWN, "boom")
            (workdir / "output.md").write_text(
                "# Out\n\n## Abstract\nx\n\n## Introduction\ny\n\n"
                "## Conclusion\nz\n\n## References\n[1] a\n",
                encoding="utf-8",
            )
            return "ok"

    for client in ("Vani", "Zeta"):
        write_paper(
            cfg.research_papers_root / MONTH / "Manoj Paper" / client, f"{client}.docx"
        )

    import paper_automation.phases as phases_module

    original_builder = phases_module.build_providers
    phases_module.build_providers = lambda _cfg: (FailsOnVani(), FailsOnVani())
    try:
        exit_code = run_pipeline(cfg, storage)
    finally:
        phases_module.build_providers = original_builder

    vani = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    zeta = cfg.research_papers_root / MONTH / "Manoj Paper" / "Zeta"
    assert exit_code == 1
    assert not (vani / "Vani_review.docx").exists()
    assert (zeta / "Correct_Zeta_paper.docx").exists()


def test_usage_limit_stops_the_phase_without_touching_later_clients(cfg, storage):
    class OutOfQuota(CliProvider):
        name = model_label = "limited"

        def _invoke(self, workdir: Path, prompt: str) -> str:
            raise ProviderError(FailureKind.USAGE_LIMIT, "usage limit reached")

    for client in ("Alpha", "Beta", "Gamma"):
        write_paper(
            cfg.research_papers_root / MONTH / "Manoj Paper" / client, f"{client}.docx"
        )

    import paper_automation.phases as phases_module

    original_builder = phases_module.build_providers
    phases_module.build_providers = lambda _cfg: (OutOfQuota(), OutOfQuota())
    try:
        run_pipeline(cfg, storage, phase="review")
    finally:
        phases_module.build_providers = original_builder

    root = cfg.research_papers_root / MONTH / "Manoj Paper"
    for client in ("Alpha", "Beta", "Gamma"):
        assert [p.name for p in (root / client).iterdir()] == [f"{client}.docx"]


def test_revision_phase_alone_skips_when_no_review_exists(cfg, storage):
    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    write_paper(folder, "Vani.docx")

    run_pipeline(cfg, storage, phase="revise")

    assert [p.name for p in folder.iterdir()] == ["Vani.docx"]


def test_validation_failure_flags_for_human_review(cfg, storage):
    """A fabricated number must produce a file, but never a COMPLETED state."""

    class Inflates(CliProvider):
        name = model_label = "inflating"

        def _invoke(self, workdir: Path, prompt: str) -> str:
            manuscript = (workdir / "manuscript.md").read_text(encoding="utf-8")
            if (workdir / "review.md").exists():
                manuscript = manuscript.replace("94.20%", "99.90%")
            (workdir / "output.md").write_text(manuscript, encoding="utf-8")
            return "ok"

    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    write_paper(folder, "Vani.docx")

    import paper_automation.phases as phases_module
    from paper_automation.state import StateStore

    original_builder = phases_module.build_providers
    phases_module.build_providers = lambda _cfg: (Inflates(), Inflates())
    try:
        run_pipeline(cfg, storage)
    finally:
        phases_module.build_providers = original_builder

    assert (folder / "Correct_Vani_paper.docx").exists()
    with StateStore(cfg.state_db) as store:
        row = store.get(MONTH, "Manoj Paper", "Vani")
    assert row["state"] == ProcessingState.REQUIRES_HUMAN_REVIEW.value
    assert "99.90" in row["reason"]


def test_original_modified_during_review_is_flagged_not_blocked(cfg, storage):
    """The doc's guard (spec section 14): the job still finishes using the version it
    captured, but the reason records that the original changed mid-run."""

    class EditsOriginalMidRun(CliProvider):
        name = model_label = "editor"

        def _invoke(self, workdir: Path, prompt: str) -> str:
            manuscript = (workdir / "manuscript.md").read_text(encoding="utf-8")
            original.write_text("The writer changed this while it was processing.")
            (workdir / "output.md").write_text(manuscript, encoding="utf-8")
            return "ok"

    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    original = write_paper(folder, "Vani.docx")

    import paper_automation.phases as phases_module
    from paper_automation.state import StateStore

    original_builder = phases_module.build_providers
    phases_module.build_providers = lambda _cfg: (EditsOriginalMidRun(), EditsOriginalMidRun())
    try:
        run_pipeline(cfg, storage, phase="review")
    finally:
        phases_module.build_providers = original_builder

    with StateStore(cfg.state_db) as store:
        row = store.get(MONTH, "Manoj Paper", "Vani")
        latest = store.latest_version(MONTH, "Manoj Paper", "Vani", "original")

    assert (folder / "Vani_review.docx").exists()
    assert "modified during processing" in row["reason"]
    assert latest["version_number"] == 2


def test_unmodified_original_leaves_no_stray_note(cfg, storage):
    from paper_automation.state import StateStore

    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    write_paper(folder, "Vani.docx")

    run_pipeline(cfg, storage, phase="review")

    with StateStore(cfg.state_db) as store:
        row = store.get(MONTH, "Manoj Paper", "Vani")
        latest = store.latest_version(MONTH, "Manoj Paper", "Vani", "original")

    assert row["reason"] == ""
    assert latest["version_number"] == 1


def test_a_full_run_leaves_finished_job_records(cfg, storage):
    from paper_automation.state import StateStore

    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    write_paper(folder, "Vani.docx")

    run_pipeline(cfg, storage)

    with StateStore(cfg.state_db) as store:
        jobs = store.list_jobs()

    phases_seen = {row["phase"] for row in jobs}
    assert phases_seen == {"review", "revise"}
    assert all(row["status"] in ("REVIEW_COMPLETED", "COMPLETED") for row in jobs)
    assert all(row["finished_at"] for row in jobs)


def test_a_stale_in_progress_job_is_requeued_at_run_start(cfg, storage):
    from paper_automation.models import Phase as PhaseEnum
    from paper_automation.state import StateStore

    folder = cfg.research_papers_root / MONTH / "Manoj Paper" / "Vani"
    write_paper(folder, "Vani.docx")

    with StateStore(cfg.state_db) as store:
        stale = store.enqueue_job(MONTH, "Manoj Paper", "Ghost", PhaseEnum.REVIEW)
        store.start_job(stale)

    run_pipeline(cfg, storage, phase="review")

    with StateStore(cfg.state_db) as store:
        row = store.get_job(stale)
    assert row["status"] == "QUEUED"
    assert "interrupted" in row["reason"]


# --- concurrent worker pool (spec section 19) ---------------------------------


def test_concurrent_workers_actually_overlap_in_time(cfg, storage):
    """Proves max_concurrent_jobs > 1 genuinely runs subprocesses in parallel,
    not just a relabelled sequential loop."""
    import threading
    import time as time_module

    lock = threading.Lock()
    counters = {"current": 0, "peak": 0}

    class Overlapping(CliProvider):
        name = model_label = "overlapping"

        def _invoke(self, workdir: Path, prompt: str) -> str:
            with lock:
                counters["current"] += 1
                counters["peak"] = max(counters["peak"], counters["current"])
            time_module.sleep(0.2)
            with lock:
                counters["current"] -= 1
            (workdir / "output.md").write_text(
                "# Out\n\n## Abstract\nx\n\n## Introduction\ny\n\n"
                "## Conclusion\nz\n\n## References\n[1] a\n",
                encoding="utf-8",
            )
            return "ok"

    for client in ("Alpha", "Beta", "Gamma"):
        write_paper(
            cfg.research_papers_root / MONTH / "Manoj Paper" / client, f"{client}.docx"
        )
    cfg.max_concurrent_jobs = 3

    import paper_automation.phases as phases_module

    original_builder = phases_module.build_providers
    phases_module.build_providers = lambda _cfg: (Overlapping(), Overlapping())
    try:
        run_pipeline(cfg, storage, phase="review")
    finally:
        phases_module.build_providers = original_builder

    assert counters["peak"] >= 2


def test_concurrent_processing_does_not_mix_up_papers(cfg, storage):
    """Each worker must get its own provider instance and its own scratch dir —
    proves that sharing isn't accidentally leaking one client's content into
    another's output under concurrency."""
    from paper_automation import docx_io

    papers = {
        "Alpha": "Alpha-only marker one one one",
        "Beta": "Beta-only marker two two two",
        "Gamma": "Gamma-only marker three three three",
    }
    for name, marker in papers.items():
        folder = cfg.research_papers_root / MONTH / "Manoj Paper" / name
        folder.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading(name, 0)
        document.add_paragraph(marker)
        document.save(folder / f"{name}.docx")
    cfg.max_concurrent_jobs = 3

    run_pipeline(cfg, storage)

    for name, marker in papers.items():
        final = cfg.research_papers_root / MONTH / "Manoj Paper" / name / f"Correct_{name}_paper.docx"
        assert final.exists()
        text = docx_io.extract(final)
        assert marker in text
        for other_name, other_marker in papers.items():
            if other_name != name:
                assert other_marker not in text


def test_concurrent_run_still_finishes_all_reviews_before_any_revision(cfg, storage):
    order: list[str] = []
    lock = __import__("threading").Lock()

    class Recording(CliProvider):
        name = model_label = "recording"

        def _invoke(self, workdir: Path, prompt: str) -> str:
            phase = "revise" if (workdir / "review.md").exists() else "review"
            with lock:
                order.append(phase)
            (workdir / "output.md").write_text(
                "# Out\n\n## Abstract\nx\n\n## Introduction\ny\n\n"
                "## Conclusion\nz\n\n## References\n[1] a\n",
                encoding="utf-8",
            )
            return "ok"

    for client in ("Vani", "R Ashok", "Jyoti R"):
        write_paper(
            cfg.research_papers_root / MONTH / "Manoj Paper" / client, f"{client}.docx"
        )
    cfg.max_concurrent_jobs = 3

    import paper_automation.phases as phases_module

    original_builder = phases_module.build_providers
    phases_module.build_providers = lambda _cfg: (Recording(), Recording())
    try:
        run_pipeline(cfg, storage)
    finally:
        phases_module.build_providers = original_builder

    assert order.count("review") == 3
    assert order.count("revise") == 3
    assert order.index("revise") > order.index("review")
    assert order[:3] == ["review", "review", "review"]


def test_each_concurrent_client_gets_its_own_job_record(cfg, storage):
    from paper_automation.state import StateStore

    for client in ("Alpha", "Beta", "Gamma"):
        write_paper(
            cfg.research_papers_root / MONTH / "Manoj Paper" / client, f"{client}.docx"
        )
    cfg.max_concurrent_jobs = 3

    run_pipeline(cfg, storage, phase="review")

    with StateStore(cfg.state_db) as store:
        jobs = store.list_jobs()

    job_ids = [row["job_id"] for row in jobs]
    assert len(job_ids) == len(set(job_ids))  # every job got a distinct id
    assert len(jobs) == 3
    assert all(row["status"] == "REVIEW_COMPLETED" for row in jobs)


def test_config_with_a_utf8_bom_loads(tmp_path):
    """Notepad and PowerShell's Out-File both write a BOM; tomllib rejects it."""
    from paper_automation import config as config_module

    path = tmp_path / "config.toml"
    path.write_text(
        '\ufeffresearch_papers_root = "%s"\n' % tmp_path.as_posix(),
        encoding="utf-8",
    )
    cfg = config_module.load(path, base_dir=tmp_path)
    assert cfg.research_papers_root == tmp_path


def test_malformed_config_reports_the_file(tmp_path):
    from paper_automation import config as config_module

    path = tmp_path / "config.toml"
    path.write_text("this is not toml at all\n", encoding="utf-8")
    with pytest.raises(config_module.ConfigError, match="not valid TOML"):
        config_module.load(path, base_dir=tmp_path)


def _write_config(tmp_path, extra=""):
    path = tmp_path / "config.toml"
    path.write_text(
        'research_papers_root = "%s"\n%s' % (tmp_path.as_posix(), extra),
        encoding="utf-8",
    )
    return path


def test_storage_backend_defaults_to_local_when_absent(tmp_path):
    from paper_automation import config as config_module

    cfg = config_module.load(_write_config(tmp_path), base_dir=tmp_path)
    assert cfg.storage_backend == "local"


def test_storage_backend_local_loads_explicitly(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(tmp_path, 'storage_backend = "local"\n')
    cfg = config_module.load(path, base_dir=tmp_path)
    assert cfg.storage_backend == "local"


def test_storage_backend_gdrive_is_rejected(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(tmp_path, 'storage_backend = "gdrive"\n')
    with pytest.raises(config_module.ConfigError, match="reserved for a future release"):
        config_module.load(path, base_dir=tmp_path)


def test_storage_backend_unknown_value_is_rejected(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(tmp_path, 'storage_backend = "dropbox"\n')
    with pytest.raises(config_module.ConfigError, match="storage_backend must be one of"):
        config_module.load(path, base_dir=tmp_path)


def test_build_storage_returns_local_storage(cfg):
    from paper_automation.storage import LocalStorage, build_storage

    assert isinstance(build_storage(cfg), LocalStorage)


def test_max_concurrent_jobs_defaults_to_one(tmp_path):
    from paper_automation import config as config_module

    cfg = config_module.load(_write_config(tmp_path), base_dir=tmp_path)
    assert cfg.max_concurrent_jobs == 1


def test_max_concurrent_jobs_reads_a_configured_value(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(tmp_path, "max_concurrent_jobs = 4\n")
    cfg = config_module.load(path, base_dir=tmp_path)
    assert cfg.max_concurrent_jobs == 4


def test_max_concurrent_jobs_rejects_zero(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(tmp_path, "max_concurrent_jobs = 0\n")
    with pytest.raises(config_module.ConfigError, match="at least 1"):
        config_module.load(path, base_dir=tmp_path)


def test_max_concurrent_jobs_rejects_a_non_integer(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(tmp_path, 'max_concurrent_jobs = "many"\n')
    with pytest.raises(config_module.ConfigError, match="whole number"):
        config_module.load(path, base_dir=tmp_path)


def test_schedule_defaults_to_nine_to_six(tmp_path):
    from datetime import time
    from paper_automation import config as config_module

    cfg = config_module.load(_write_config(tmp_path), base_dir=tmp_path)
    assert cfg.schedule_start == time(9, 0)
    assert cfg.schedule_end == time(18, 0)
    assert cfg.scan_interval_minutes == 15


def test_schedule_reads_configured_values(tmp_path):
    from datetime import time
    from paper_automation import config as config_module

    path = _write_config(
        tmp_path,
        'schedule_start = "08:30"\nschedule_end = "17:45"\nscan_interval_minutes = 10\n',
    )
    cfg = config_module.load(path, base_dir=tmp_path)
    assert cfg.schedule_start == time(8, 30)
    assert cfg.schedule_end == time(17, 45)
    assert cfg.scan_interval_minutes == 10


def test_schedule_rejects_a_malformed_time(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(tmp_path, 'schedule_start = "not-a-time"\n')
    with pytest.raises(config_module.ConfigError, match="HH:MM format"):
        config_module.load(path, base_dir=tmp_path)


def test_schedule_rejects_start_at_or_after_end(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(
        tmp_path, 'schedule_start = "18:00"\nschedule_end = "09:00"\n'
    )
    with pytest.raises(config_module.ConfigError, match="must be before"):
        config_module.load(path, base_dir=tmp_path)


def test_scan_interval_rejects_zero(tmp_path):
    from paper_automation import config as config_module

    path = _write_config(tmp_path, "scan_interval_minutes = 0\n")
    with pytest.raises(config_module.ConfigError, match="at least 1"):
        config_module.load(path, base_dir=tmp_path)


def test_build_storage_rejects_unsupported_backend(cfg):
    from paper_automation import config as config_module
    from paper_automation.storage import build_storage

    cfg.storage_backend = "dropbox"
    with pytest.raises(config_module.ConfigError, match="Unsupported storage_backend"):
        build_storage(cfg)
