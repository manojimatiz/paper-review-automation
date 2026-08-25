import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402

from paper_automation.config import Config  # noqa: E402
from paper_automation.storage import LocalStorage  # noqa: E402


@pytest.fixture
def storage() -> LocalStorage:
    return LocalStorage()


@pytest.fixture
def cfg(tmp_path: Path) -> Config:
    return Config(
        research_papers_root=tmp_path / "Research Papers",
        timezone="Asia/Kolkata",
        state_db=tmp_path / "state" / "processing.sqlite3",
        log_dir=tmp_path / "logs",
        scratch_dir=tmp_path / "scratch",
    )


@pytest.fixture
def make_client(cfg: Config):
    """Create a client folder containing the named files, and return its path."""

    def _make(
        *files: str,
        month: str = "August 2026",
        employee: str = "Manoj Paper",
        client: str = "Vani",
        subdirs: tuple[str, ...] = (),
    ) -> Path:
        folder = cfg.research_papers_root / month / employee / client
        folder.mkdir(parents=True, exist_ok=True)
        for name in files:
            if name.lower().endswith(".docx") and not name.startswith("~$"):
                # A real document, so tests that actually open the file work.
                document = Document()
                document.add_heading("Test Manuscript", 0)
                document.add_paragraph(
                    "We fine-tuned ResNet-50 and reached 94.20% accuracy "
                    "with an F1 of 0.912 on the evaluation set."
                )
                document.save(folder / name)
            else:
                (folder / name).write_text("content", encoding="utf-8")
        for name in subdirs:
            (folder / name).mkdir(exist_ok=True)
        return folder

    return _make
