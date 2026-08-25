"""Round-tripping .docx through markdown."""

from docx import Document

from paper_automation import docx_io
from paper_automation.docx_io import DocxError

import pytest


def make_paper(path, *, heading="Introduction", body="Some prose.") -> None:
    document = Document()
    document.add_heading("A Study of Things", level=1)
    document.add_heading(heading, level=2)
    document.add_paragraph(body)
    document.add_paragraph("First point", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Model"
    table.cell(0, 1).text = "Accuracy"
    table.cell(1, 0).text = "ResNet"
    table.cell(1, 1).text = "94.20%"
    document.save(str(path))


def test_extract_preserves_structure(tmp_path):
    path = tmp_path / "paper.docx"
    make_paper(path)

    markdown = docx_io.extract(path)

    assert "# A Study of Things" in markdown
    assert "## Introduction" in markdown
    assert "- First point" in markdown
    assert "| Model | Accuracy |" in markdown
    assert "94.20%" in markdown


def test_build_renders_markdown(tmp_path):
    target = tmp_path / "out.docx"
    markdown = (
        "# Title\n\n"
        "## Section\n\n"
        "Body text with **bold** and *italic*.\n\n"
        "- bullet one\n- bullet two\n\n"
        "1. step one\n\n"
        "| A | B |\n|---|---|\n| 1 | 2 |\n"
    )

    docx_io.build(markdown, target)

    document = Document(str(target))
    texts = [p.text for p in document.paragraphs]
    assert "Title" in texts
    assert "Section" in texts
    assert any("bold" in t and "italic" in t for t in texts)
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "A"


def test_build_refuses_to_overwrite(tmp_path):
    target = tmp_path / "out.docx"
    target.write_text("existing", encoding="utf-8")

    with pytest.raises(DocxError):
        docx_io.build("# New", target)

    assert target.read_text(encoding="utf-8") == "existing"


def test_round_trip_keeps_numbers_and_names(tmp_path):
    source = tmp_path / "paper.docx"
    make_paper(source, body="ResNet reached an accuracy of 94.20% on CIFAR-10.")
    target = tmp_path / "rebuilt.docx"

    markdown = docx_io.extract(source)
    docx_io.build(markdown, target)
    round_tripped = docx_io.extract(target)

    assert "94.20%" in round_tripped
    assert "CIFAR-10" in round_tripped
    assert "ResNet" in round_tripped


def test_extract_rejects_a_non_docx(tmp_path):
    path = tmp_path / "fake.docx"
    path.write_text("not a real docx", encoding="utf-8")

    with pytest.raises(DocxError):
        docx_io.extract(path)


def test_is_readable(tmp_path):
    good, bad = tmp_path / "g.docx", tmp_path / "b.docx"
    make_paper(good)
    bad.write_text("nope", encoding="utf-8")

    assert docx_io.is_readable(good)
    assert not docx_io.is_readable(bad)
